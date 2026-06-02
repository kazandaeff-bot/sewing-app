import sys, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Noto Sans SC'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.3

# Title
title = doc.add_heading('Ревью приложения: Швейное производство', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Date
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Дата: 30.05.2026')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# ============ SUMMARY ============
doc.add_heading('Общее резюме', level=1)
p = doc.add_paragraph()
p.add_run('Приложение представляет собой полнофункциональную платформу автоматизации швейного производства на стеке Next.js 16 + Prisma + SQLite + Bun. Реализован полный производственный цикл: Планы пошива → Раскрой → Задания швеям → Пошив → ВТО (утюжка) → ОТК (контроль качества) → Распределение по городам → Упаковка в короба → Отгрузка. Также есть модули: справочники, материалы, зарплата, печать документов.').font.size = Pt(11)

p = doc.add_paragraph()
p.add_run('Оценка: ').bold = True
p.add_run('Проект находится на продвинутой стадии (70-75% готовности). Основной функционал работает, но есть критические баги, дыры в безопасности и архитектурные проблемы, которые нужно решить перед боевым запуском.')

# ============ WHAT WORKS ============
doc.add_heading('Что работает', level=1)

items_work = [
    ('Авторизация и роли', 'Реализована система ролей (supervisor, sewer, qc, ironing, cutter, seller, technologist, customer). bcrypt хеширование паролей, JWT-подобные токены в httpOnly cookie, withAuth() middleware для API маршрутов.'),
    ('Планы пошива', 'Полный CRUD: создание, редактирование, утверждение, дополнение (supplement), перевод в работу/отгрузку. Приоритеты (срочный/обычный/низкий), дедлайны, прогресс-бар по позициям.'),
    ('Раскрой', 'Автосоздание при утверждении плана, ввод фактических количеств, автосписание материалов, автоматические остатки кроя.'),
    ('Задания швеям', 'Распределение позиций раскроя по швеям, контроль превышения количества, печать нарядов.'),
    ('Пошив (швея)', 'Швея видит свои задания, может начать работу, отшить полностью или частями, отправить на ВТО.'),
    ('ВТО (утюжка)', 'Отдельная вкладка для утюжки, группировка по заданиям, отметка «отглажено», расчёт зарплаты утюжки.'),
    ('ОТК (контроль качества)', 'Поштучная приёмка/возврат на переделку, создание переделок с причинами, фильтрация по статусу, расчёт зарплаты ОТК.'),
    ('Распределение по городам', 'Выбор плана, добавление городов, распределение количеств, прогресс-бар, утверждение/отметка распределения.'),
    ('Короба', 'Генерация из утверждённого плана распределения, ввод фактических количеств, продвижение статуса (формирование → собран → проверен → отгружен), печать.'),
    ('Справочники', 'Изделия (со ставками, размерами, цветами, комплектами), сотрудники, города, типы коробов, заказчики, типы материалов, нормы расхода, приход/расход материалов.'),
    ('Зарплата', 'Общий хук useSalaryCalculation для швей, ОТК и утюжки. Расшифровка по изделиям, фильтр по периоду (неделя/месяц/всё время).'),
    ('Печать документов', 'Наряды раскроя, задания швеям, упаковочные листы — генерация HTML для печати.'),
    ('Роль заказчика', 'Заказчик видит только свои планы, распределения, короба и материальный баланс.'),
]

for title, desc in items_work:
    p = doc.add_paragraph()
    p.add_run(f'✅ {title}: ').bold = True
    p.add_run(desc)

# ============ CRITICAL BUGS FIXED ============
doc.add_heading('Критические баги (ИСПРАВЛЕНО)', level=1)

p = doc.add_paragraph()
p.add_run('plans.filter is not a function').bold = True
p.add_run(' — когда API возвращает {error: "Не авторизован"} вместо массива, React Query прокидывает этот объект в компонент, и вызов .filter() на нём падает. Эта ошибка была на скриншоте пользователя в city-distribution-tab.tsx. Исправлено добавлением Array.isArray() проверок во ВСЕ компоненты, использующие raw fetch:')

bug_items = [
    'city-distribution-tab.tsx — plans, sellerPlans, cities',
    'boxes-tab.tsx — boxes, sellerPlans',
    'sewing-tasks-tab.tsx — sewingTasks, cuttingPlans, employees',
    'cutting-plans-tab.tsx — cuttingPlans',
    'cutting-leftovers-tab.tsx — leftovers, customers',
    'employees-tab.tsx — employees, customers',
    'products-tab.tsx — products',
]

for item in bug_items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
p.add_run('Паттерн: ').bold = True
p.add_run('const data = Array.isArray(rawData) ? rawData : [] — применён везде вместо прямого использования данных из useQuery.')

# ============ REMAINING BUGS ============
doc.add_heading('Оставшиеся баги и проблемы', level=2)

doc.add_heading('🔴 Критические', level=3)

critical = [
    ('Токен не верифицируется (auth.ts)', 'verifyToken() только декодирует base64 тело, но НЕ проверяет подпись. Любой может подделать токен, закодировав произвольный JSON в base64. Нужно использовать настоящую HMAC-SHA256 верификацию или полноценный JWT (jsonwebtoken).'),
    ('Seed-маршрут без авторизации', 'GET /api/seed?force=true доступен без авторизации — любой может удалить все данные. Нужно добавить withAuth(seedHandler, ["supervisor"]).'),
    ('Нет Error Boundary', 'Если любой компонент выбрасывает исключение, всё приложение падает с белым экраном. Нужно добавить React Error Boundary на уровне страницы и табов.'),
    ('payReworkToQC не в типах', 'Product интерфейс в types/index.ts не содержит поле payReworkToQC, хотя оно есть в Prisma-схеме и API. Это может привести к ошибкам TypeScript и некорректной работе.'),
]

for title, desc in critical:
    p = doc.add_paragraph()
    p.add_run(f'🔴 {title}: ').bold = True
    p.add_run(desc)

doc.add_heading('🟠 Серьёзные', level=3)

serious = [
    ('Неиспользуемый api-client.ts', 'Создан общий API клиент (apiGet/apiPost/apiPatch/apiDelete) с правильной обработкой ошибок, но 90% компонентов используют сырой fetch().then(r => r.json()) без проверки res.ok. Нужно мигрировать все компоненты на api-client.'),
    ('Дублирующиеся маршруты', '/api/auth/session дублирует /api/auth/me, /api/hello дублирует /api/. Нужно удалить дубликаты.'),
    ('Роль seller перенаправляет на заглушку', 'Продавец видит «Перейдите в раздел Производство», но кнопка просто перезагружает страницу. Нужно дать продавцу доступ к вкладкам «Города» и «Короба».'),
    ('Роли technologist и cutter — заглушки', 'Технолог и раскройщик видят «Раздел в разработке». Cutter должен видеть раскрой, technologist — изделия и нормы.'),
    ('Мутации без проверки res.ok', 'Большинство useMutation используют .then(r => r.json()) без проверки r.ok. При ошибке сервера (401/403/500) мутация «успешна», но данные — объект ошибки. Нужно добавить проверку во все мутации.'),
    ('QC зарплата не учитывает payReworkToQC', 'Когда payReworkToQC=true, проверка переделок ОТК должна оплачиваться дополнительно, но текущий расчёт показывает «без доплаты за переделки» в любом случае.'),
]

for title, desc in serious:
    p = doc.add_paragraph()
    p.add_run(f'🟠 {title}: ').bold = True
    p.add_run(desc)

doc.add_heading('🟡 Средние', level=3)

medium = [
    ('Много any в коде', 'SewerTab, QCTab и другие компоненты используют task: any, item: any. Нужно заменить на конкретные типы из @/types.'),
    ('Нет глобальной обработки 401', 'При истечении токена API возвращает 401, но фронтенд не перенаправляет на страницу логина автоматически. Нужно добавить перехватчик 401 в api-client.'),
    ('Отсутствует валидация на фронтенде', 'Нет клиентской валидации форм (только серверная Zod). Нужно добавить react-hook-form + zod для форм создания/редактирования.'),
    ('Нет пагинации', 'Все запросы возвращают полные списки. При росте данных (1000+ записей) приложение станет медленным.'),
    ('Нет оптимистичных обновлений', 'Каждая мутация ждёт ответа сервера. Для UX лучше использовать optimistic updates в React Query.'),
]

for title, desc in medium:
    p = doc.add_paragraph()
    p.add_run(f'🟡 {title}: ').bold = True
    p.add_run(desc)

# ============ SECURITY ============
doc.add_heading('Безопасность', level=1)

security_items = [
    ('Токены', 'Текущая реализация — base64-encoded JSON без криптографической верификации. Уязвима к подделке. Решение: перейти на jsonwebtoken с HMAC-SHA256 или RS256.'),
    ('JWT_SECRET', 'Хардкод "dev-secret-change-me". В проде MUST быть через env-переменную с длинным случайным значением.'),
    ('CORS', 'Middleware устанавливает no-cache заголовки, но нет явного контроля CORS. В проде нужно ограничить разрешённые origins.'),
    ('Rate limiting', 'Нет ограничений на количество запросов, особенно на /api/auth/login. Уязвимо к брутфорсу.'),
    ('SQL injection', 'Prisma параметризует запросы — защита есть. Но raw queries (если появятся) нужно проверять отдельно.'),
    ('XSS', 'React по умолчанию экранирует. Но dangerouslySetInnerHTML или innerHTML (в печати) — потенциальные векторы.'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Аспект'
hdr[1].text = 'Текущее состояние'
hdr[2].text = 'Рекомендация'

for aspect, current, rec in [
    ('Токены', 'Base64 без верификации', 'JWT с HMAC-SHA256'),
    ('JWT_SECRET', 'Хардкод', 'Env-переменная'),
    ('CORS', 'Нет ограничений', 'Ограничить origins'),
    ('Rate limiting', 'Нет', 'Добавить на /api/auth/login'),
    ('SQL injection', 'Защищено (Prisma)', 'Проверять raw queries'),
    ('XSS', 'Частично', 'Очистить HTML в печати'),
]:
    row = table.add_row().cells
    row[0].text = aspect
    row[1].text = current
    row[2].text = rec

# ============ ARCHITECTURE ============
doc.add_heading('Архитектура', level=1)

arch_items = [
    ('Плюсы', [
        'Разделение на табы-компоненты — хорошая модульность',
        'Общие хуки (useSalaryCalculation, useItemRows, useProductColorSelect) — устраняют дублирование',
        'Zod-схемы с validateBody на всех API маршрутах — надёжная валидация',
        'Prisma $transaction() на критичных маршрутах (3 маршрута) — атомарность',
        'withAuth() middleware — централизованная авторизация',
    ]),
    ('Минусы', [
        'production-tabs.tsx — файл-прокладка, все импорты идут напрямую из tabs/',
        'Крупные компоненты (references-tab ~1000 строк, sewing-plans-tab ~1000 строк) — нужно разбить',
        'Смешивание стилей: часть компонентов использует apiGet, часть — raw fetch',
        'Нет state management для cross-tab данных (React Query кэш частично решает)',
        'Сид-данные создают бенчмарк-записи, но не покрывают все бизнес-сценарии',
    ]),
]

for category, items in arch_items:
    doc.add_heading(category, level=2)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# ============ RECOMMENDATIONS ============
doc.add_heading('План улучшений (приоритеты)', level=1)

recs = [
    ('1', '🔴 Починить верификацию токена', 'Заменить базовый base64 «JWT» на настоящий JWT с подписью. Использовать jsonwebtoken или jose. Это блокер для продакшена.', '1-2 дня'),
    ('2', '🔴 Закрыть seed-маршрут авторизацией', 'Обернуть GET /api/seed в withAuth(handler, ["supervisor"]).', '30 минут'),
    ('3', '🔴 Добавить Error Boundary', 'Обернуть приложение в React Error Boundary с понятным UI ошибки и кнопкой «Перезагрузить».', '2-3 часа'),
    ('4', '🟠 Мигрировать на api-client', 'Заменить все сырые fetch() на apiGet/apiPost/apiPatch/apiDelete. Это унифицирует обработку ошибок и добавит автоматический перехват 401.', '1 день'),
    ('5', '🟠 Дать продавцу доступ к вкладкам', 'Seller должен видеть «Города» и «Короба» вместо заглушки.', '2-3 часа'),
    ('6', '🟠 Починить QC зарплату с payReworkToQC', 'Когда payReworkToQC=true, добавить reworkRate × количество переделок к зарплате ОТК.', '2-3 часа'),
    ('7', '🟡 Убрать any-типы', 'Заменить task: any, item: any на конкретные типы из @/types.', '1 день'),
    ('8', '🟡 Добавить автологаут при 401', 'В api-client при получении 401 — очистить cookie и перенаправить на /login.', '2 часа'),
    ('9', '🟡 Реализовать Distribution/Allocation', 'Запрошенная фича: равное распределение по городам, автоназначение раскроя швеям с ограничениями (1 цвет на швею, смешанные размеры, кратность пачкам).', '3-5 дней'),
    ('10', '🔵 Юнит-тесты', 'Покрыть критичные API маршруты и хуки тестами. Приоритет: auth, sewing-tasks, cutting-plans, useSalaryCalculation.', '3-5 дней'),
]

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
hdr2[0].text = '#'
hdr2[1].text = 'Задача'
hdr2[2].text = 'Описание'
hdr2[3].text = 'Срок'

for num, title, desc, time in recs:
    row = table2.add_row().cells
    row[0].text = num
    row[1].text = title
    row[2].text = desc
    row[3].text = time

# ============ CONCLUSION ============
doc.add_heading('Итог', level=1)

p = doc.add_paragraph()
p.add_run('Проект находится на хорошей стадии. Основной производственный цикл реализован и функционален. Ключевые проблемы — безопасность (токены без верификации), отсутствие Error Boundary и необработанные ошибки API на фронтенде — были частично исправлены в этом ревью (Array.isArray проверки добавлены во все 7 компонентов). Оставшиеся проблемы имеют чёткий план исправления.')

p = doc.add_paragraph()
p.add_run('Рекомендация: ').bold = True
p.add_run('Перед запуском в продакшен обязательно выполнить пункты 1-6 из плана улучшений. Пункты 7-10 можно делать параллельно с разработкой новых фич.')

doc.save('/home/z/my-project/download/mentor_review.docx')
print('Done: /home/z/my-project/download/mentor_review.docx')
